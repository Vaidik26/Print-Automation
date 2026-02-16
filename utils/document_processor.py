"""
Document Processor Module
Handles Word document template parsing and content replacement.
"""

import re
import copy
from docx import Document
from docx.shared import Pt
from io import BytesIO
from typing import List, Set, Dict, Tuple


class DocumentProcessor:
    """Process Word documents with placeholder replacement."""

    # Pattern to match placeholders like {column_name}, {Column Name}, {COLUMN_NAME}
    PLACEHOLDER_PATTERN = re.compile(r"\{([^}]+)\}")

    def __init__(self, template_bytes: bytes):
        """
        Initialize with template document bytes.

        Args:
            template_bytes: The Word document as bytes
        """
        self.template_bytes = template_bytes
        self.template_doc = Document(BytesIO(template_bytes))
        self.placeholders = self._extract_placeholders()

    def _extract_text_from_paragraph(self, paragraph) -> str:
        """Extract full text from a paragraph, handling split runs."""
        return paragraph.text

    def _extract_placeholders_from_text(self, text: str) -> Set[str]:
        """Extract all placeholders from a text string."""
        matches = self.PLACEHOLDER_PATTERN.findall(text)
        return set(m.strip() for m in matches)

    def _extract_placeholders(self) -> Set[str]:
        """
        Extract all unique placeholders from the template.
        Searches in paragraphs, tables, headers, and footers.

        Returns:
            Set of placeholder names (without braces)
        """
        placeholders = set()

        # Extract from main document paragraphs
        for paragraph in self.template_doc.paragraphs:
            text = self._extract_text_from_paragraph(paragraph)
            placeholders.update(self._extract_placeholders_from_text(text))

        # Extract from tables
        for table in self.template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = self._extract_text_from_paragraph(paragraph)
                        placeholders.update(self._extract_placeholders_from_text(text))

        # Extract from headers and footers
        for section in self.template_doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = self._extract_text_from_paragraph(paragraph)
                    placeholders.update(self._extract_placeholders_from_text(text))
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = self._extract_text_from_paragraph(paragraph)
                                placeholders.update(
                                    self._extract_placeholders_from_text(text)
                                )

            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = self._extract_text_from_paragraph(paragraph)
                    placeholders.update(self._extract_placeholders_from_text(text))
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = self._extract_text_from_paragraph(paragraph)
                                placeholders.update(
                                    self._extract_placeholders_from_text(text)
                                )

        return placeholders

    def get_placeholders(self) -> List[str]:
        """
        Get list of all placeholders found in the template.

        Returns:
            Sorted list of placeholder names
        """
        return sorted(list(self.placeholders))

    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """
        Replace placeholders in a paragraph while preserving formatting.

        This handles cases where placeholders might be split across multiple runs.
        """
        # Get the full text
        full_text = paragraph.text

        # Check if there are any placeholders to replace
        if not self.PLACEHOLDER_PATTERN.search(full_text):
            return

        # Replace all placeholders in the full text
        new_text = full_text
        for placeholder, value in replacements.items():
            # Create pattern to match {Placeholder}, { Placeholder }, {PLACEHOLDER}, etc.
            # We use the key from replacements which is already stripped
            pattern = r"\{\s*" + re.escape(placeholder) + r"\s*\}"
            new_text = re.sub(
                pattern, 
                str(value) if value is not None else "",
                new_text,
                flags=re.IGNORECASE 
            )

        # If text changed, we need to update the paragraph
        if new_text != full_text:
            # Store the formatting of the first run (if any)
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                # Set new text to first run
                first_run.text = new_text
            else:
                # No runs, just set text directly
                paragraph.text = new_text

    def _replace_in_table(self, table, replacements: Dict[str, str]):
        """Replace placeholders in a table."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)
                # Handle nested tables
                for nested_table in cell.tables:
                    self._replace_in_table(nested_table, replacements)

    def generate_document(
        self, data: Dict[str, str], preserved_placeholders: List[str] = None
    ) -> bytes:
        """
        Generate a new document with placeholders replaced by data values.

        Args:
            data: Dictionary mapping placeholder names to values
            preserved_placeholders: List of placeholders to keep intact (e.g. {Signature})

        Returns:
            Generated document as bytes
        """
        if preserved_placeholders is None:
            # Default reserved placeholders for DocuSign
            preserved_placeholders = ["Signature"]

        # Create a fresh copy of the template
        doc = Document(BytesIO(self.template_bytes))

        # Build replacement dictionary
        replacements = {}
        for placeholder in self.placeholders:
            # Skip if explicitly preserved
            if placeholder in preserved_placeholders:
                continue

            if placeholder in data:
                replacements[placeholder] = data[placeholder]
            else:
                # Check for case-insensitive match
                for key in data:
                    if key.lower() == placeholder.lower():
                        replacements[placeholder] = data[key]
                        break
                else:
                    replacements[placeholder] = ""  # Default to empty if not found

        # Replace in main document paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in doc.tables:
            self._replace_in_table(table, replacements)

        # Replace in headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)
                for table in section.header.tables:
                    self._replace_in_table(table, replacements)

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_text_in_paragraph(paragraph, replacements)
                for table in section.footer.tables:
                    self._replace_in_table(table, replacements)

        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def generate_documents(
        self, data_rows: List[Dict[str, str]], filename_column: str = None
    ) -> List[Tuple[str, bytes]]:
        """
        Generate multiple documents from a list of data rows.

        Args:
            data_rows: List of dictionaries, each containing data for one document
            filename_column: Column to use for naming files (optional)

        Returns:
            List of tuples (filename, document_bytes)
        """
        documents = []

        for idx, row_data in enumerate(data_rows, start=1):
            # Generate filename
            if filename_column and filename_column in row_data:
                filename = f"{row_data[filename_column]}.docx"
                # Clean filename of invalid characters
                filename = re.sub(r'[<>:"/\\|?*]', "_", filename)
            else:
                filename = f"document_{idx:04d}.docx"

            # Generate document
            doc_bytes = self.generate_document(row_data)
            documents.append((filename, doc_bytes))

        return documents
